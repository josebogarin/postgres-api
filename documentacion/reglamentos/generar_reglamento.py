# -*- coding: utf-8 -*-
# Generador del Reglamento de Clubes (Word). Parametrizado por el link de inicio.
# Uso:  python generar_reglamento.py [BASE_URL]
#   BASE_URL por defecto = https://cupped-oink-thousand.ngrok-free.dev
import sys, os
BASE_URL = (sys.argv[1] if len(sys.argv) > 1 else "https://cupped-oink-thousand.ngrok-free.dev").rstrip("/")
LINK = BASE_URL + "/becbuc-live"
_HERE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(_HERE, "..", "..", "backend", "static", "becbuc-logo.jpeg")
OUT_DOCX = os.path.join(_HERE, "BECBUC_Reglamento_Clubes_v1.docx")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN=RGBColor(0x00,0xA6,0x51); ORANGE=RGBColor(0xE0,0x50,0x20)
DARK=RGBColor(0x1a,0x1d,0x23); GREY=RGBColor(0x55,0x5b,0x66); WHITE=RGBColor(0xFF,0xFF,0xFF)

doc=Document()
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11); st.font.color.rgb=DARK

def shade(cell,hexcolor):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor); tcPr.append(sh)
def h(text,color=GREEN,size=15,sb=14,sa=4):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; return p
def body(text,size=11,italic=False,color=DARK,sa=6,bold=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.font.size=Pt(size); r.italic=italic; r.font.color.rgb=color; r.bold=bold; return p
def bullet(text,lead=None):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
    if lead:
        r=p.add_run(lead); r.bold=True; r.font.color.rgb=DARK
    r2=p.add_run(text); r2.font.color.rgb=DARK; return p
def table(headers,rows,widths,fill='00A651'):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    hdr=t.rows[0].cells
    for i,ht in enumerate(headers):
        shade(hdr[i],fill); pa=hdr[i].paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rn=pa.add_run(ht); rn.bold=True; rn.font.color.rgb=WHITE; rn.font.size=Pt(10.5); hdr[i].width=Inches(widths[i])
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].width=Inches(widths[i]); pa=cs[i].paragraphs[0]
            pa.alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT
            rn=pa.add_run(str(val)); rn.font.size=Pt(10.5)
            if i==0: rn.bold=True
    doc.add_paragraph().paragraph_format.space_after=Pt(2); return t
def add_hyperlink(paragraph,url,text,color='0563C1',size=26):
    part=paragraph.part
    r_id=part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True)
    hl=OxmlElement('w:hyperlink'); hl.set(qn('r:id'),r_id)
    rn=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    c=OxmlElement('w:color'); c.set(qn('w:val'),color); rPr.append(c)
    u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u)
    sz=OxmlElement('w:sz'); sz.set(qn('w:val'),str(size)); rPr.append(sz)
    rn.append(rPr); t=OxmlElement('w:t'); t.text=text; rn.append(t); hl.append(rn); paragraph._p.append(hl)

# ---- header/footer (no en la caratula) ----
sec=doc.sections[0]; sec.different_first_page_header_footer=True
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
hp.add_run().add_picture(LOGO,height=Inches(0.30))
hr=hp.add_run('   BECBUC · Reglamento de Apuestas de Clubes'); hr.font.size=Pt(9); hr.bold=True; hr.font.color.rgb=GREY
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run('BECBUC · Versión 1 — Agosto 2026'); fr.font.size=Pt(8); fr.font.color.rgb=GREY

# ---- CARATULA ----
for _ in range(3): doc.add_paragraph()
pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER
pic.add_run().add_picture(LOGO,width=Inches(3.3))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(20)
r=p.add_run('Reglamento de Apuestas'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=GREEN
line=doc.add_paragraph(); line.alignment=WD_ALIGN_PARAGRAPH.CENTER
pPr=line._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom')
bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'12'); bt.set(qn('w:space'),'6'); bt.set(qn('w:color'),'E05020'); pbdr.append(bt); pPr.append(pbdr)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6)
r=p.add_run('Torneos de Clubes'); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=ORANGE
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Copa Libertadores  ·  Copa Sudamericana'); r.font.size=Pt(13); r.font.color.rgb=GREY
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(34)
r=p.add_run('Guía para el apostador'); r.font.size=Pt(12); r.italic=True; r.font.color.rgb=DARK
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Versión 1 — Agosto 2026'); r.font.size=Pt(11); r.font.color.rgb=GREY
doc.add_page_break()

# ---- CONTENIDO ----
h('¿Dónde se apuesta?', ORANGE, 15, 4, 6)
body('Se apuesta en BECBUC Live, desde el celular o la computadora. Entrás al link, elegís tu nombre de la lista y cargás tus pronósticos partido por partido.')
pl=doc.add_paragraph(); pl.paragraph_format.space_after=Pt(4)
lead=pl.add_run('Link: '); lead.bold=True; lead.font.color.rgb=DARK
add_hyperlink(pl,LINK,LINK)
body('El link puede cambiar; siempre te compartimos el vigente antes de cada fecha.', italic=True, color=GREY, size=10)

h('Lo esencial en 30 segundos', ORANGE, 15, 12, 6)
bullet(' apostás partido por partido en las eliminatorias (playoffs).','Cómo se juega:')
bullet(' cada llave son dos partidos (ida y vuelta); solo la final es un partido único.','Ida y vuelta:')
bullet(' cuanto más avanza el torneo, más valen los aciertos. La final vale muchísimo.','Sube por ronda:')
bullet(' acertar quién pasa la llave, un comodín y la definición por penales pueden multiplicar tus puntos.','Multiplicadores:')
bullet(' nada está definido hasta el final. Con aciertos de mérito se puede remontar hasta el último partido.','La gracia:')

h('1. Cómo funciona')
body('Se apuesta en la fase de eliminación directa (playoffs). Cada eliminatoria (o "llave") se juega a dos partidos: uno de ida y uno de vuelta. La única excepción es la final, que es un solo partido.')
bullet(' arranca en octavos de final (los 16avos ya se jugaron y no otorgan puntos).','Copa Sudamericana:')
bullet(' arranca directamente en octavos de final.','Copa Libertadores:')
bullet(' no hay partido por el tercer puesto: de semifinales se pasa directo a la final.','Importante:')

h('2. Qué pronosticás en cada partido')
body('Para cada partido (ida y vuelta por separado) cargás tu pronóstico:')
bullet(' cuántos goles hace el equipo local y cuántos el visitante.','El marcador:')
bullet(' cuántas amarillas y cuántas rojas habrá en el partido (el total).','Tarjetas:')
bullet(' cuántos penales se cobran durante el juego (el total).','Penales en el juego:')
bullet(' cuántos cambios hace cada equipo.','Sustituciones:')
bullet(' en qué minuto se convierte el primer gol del partido.','Minuto del primer gol:')
body('Y solo en las llaves donde, según tus marcadores, la serie termina empatada y se define por penales, pronosticás también el resultado de la tanda.', italic=True, color=GREY)

h('3. Cuánto vale acertar el resultado')
body('Es el corazón del juego. En cada partido:')
bullet(' acertar quién gana (o si empatan).','Resultado:')
bullet(' acertar el marcador clavado (por ejemplo 2-1). Si acertás el marcador exacto, cobrás además los puntos del resultado.','Marcador exacto:')
table(['Fase','Acertar el resultado','Acertar el marcador exacto'],
      [['Octavos','4','8'],['Cuartos','12','24'],['Semis','30','60'],['Final','75','150']],[1.6,2.4,2.6])
body('Ejemplo: en la final pronosticás 1-0 y sale 1-0 → cobrás 75 (resultado) + 150 (marcador exacto) = 225 puntos en ese partido.', italic=True, color=GREY)

h('4. Los ítems del partido (tarjetas, penales, cambios)')
body('Estos ítems solo suman si acertás la cantidad EXACTA y hubo al menos un evento. Si el partido no tuvo (por ejemplo, cero rojas), ese ítem no suma aunque hayas puesto cero. Los puntos crecen por fase:')
table(['Ítem','Octavos','Cuartos','Semis','Final'],
      [['Rojas / Amarillas / Penales en juego','3','5','8','12'],
       ['Sustituciones (por cada equipo)','2','3','4','6']],[3.0,1.0,1.0,1.0,1.0])
h('El minuto del primer gol: un premio especial', ORANGE, 13, 8, 4)
body('Si acertás el minuto exacto del primer gol de un partido, ¡ese partido vale el DOBLE! Es un acierto muy difícil, así que cuando cae, te da un gran salto.')

h('5. Los multiplicadores: donde se gana el torneo')
h('El cruce: acertar quién pasa', GREEN, 13, 8, 4)
body('No hace falta elegir aparte quién pasa: se calcula solo con los marcadores que cargaste (sumando ida y vuelta). Si tu pronóstico deja a un equipo como clasificado y ese equipo realmente pasa, ganás el premio del cruce sobre los puntos de esa llave:')
bullet(' se multiplican por 2 los puntos de esas dos llaves.','Si acertás los dos equipos que se van a enfrentar en la ronda siguiente:')
bullet(' sumás un bono fijo por ese acierto, que crece por fase: 10 puntos en octavos, 20 en cuartos, 40 en semifinales (no se multiplica).','Si acertás solo uno de los dos:')
h('El comodín: tu jugada estratégica', GREEN, 13, 8, 4)
body('Tenés un (1) comodín para todo el torneo. Lo usás en la llave que quieras y multiplica por 3 todos los puntos de esa llave (ida + vuelta). Se puede usar de octavos a semifinales; en la final ya no. Si no lo usaste antes, se pierde. Elegí bien el momento: usado en una buena llave de semis, da un salto enorme.')
h('La definición por penales', GREEN, 13, 8, 4)
body('Si tus marcadores de ida y vuelta dejan la serie empatada (misma diferencia de gol) y por lo tanto se define por penales, podés pronosticar la tanda. Si acertás la definición, los puntos de esa llave se multiplican por 2. Ojo: solo puede cargar la tanda quien, con sus marcadores, lleva la serie a penales.')

h('6. Los premios globales (al final del torneo)')
body('Antes o durante el torneo pronosticás quién sale campeón y quién subcampeón:')
bullet(' 50 puntos.','Acertar el campeón:')
bullet(' 50 puntos.','Acertar el subcampeón:')
bullet(' el total se multiplica por 2 (hasta 200 puntos).','Si acertás el orden exacto (campeón y subcampeón en su lugar):')

h('7. Ejemplos')
h('Ejemplo 1 — Una llave de octavos', ORANGE, 12.5, 8, 3)
bullet(' pronosticaste 1-0 y salió 1-0 → acertás resultado (4) + marcador exacto (8) = 12 puntos.','Ida:')
bullet(' pronosticaste 0-2 y salió 0-1 → acertás el resultado, pero no el marcador = 4 puntos.','Vuelta:')
bullet(' 12 + 4 = 16 puntos. Y como por tus marcadores tu equipo clasifica, quedás con chances del premio del cruce.','Total de la llave:')
h('Ejemplo 2 — El comodín en semifinales', ORANGE, 12.5, 8, 3)
body('En una llave de semis sacaste 90 puntos. Le pusiste el comodín → 90 × 3 = 270 puntos de esa sola llave.')
h('Ejemplo 3 — El minuto en la final', ORANGE, 12.5, 8, 3)
body('En la final acertaste el marcador exacto (225 puntos) y además el minuto del primer gol → ese partido vale el doble: 450 puntos.')
h('Ejemplo 4 — El cruce', ORANGE, 12.5, 8, 3)
body('En octavos acertaste que pasan tanto Boca como River, y en cuartos les tocaba enfrentarse. Como acertaste los dos, los puntos que sacaste en esas dos llaves de octavos se multiplican por 2.')

h('8. Reglas de oro')
bullet(' nada se decide antes de tiempo: los puntos crecen ronda a ronda y la final vale muchísimo. Siempre hay chances de remontar.','El torneo se juega hasta el final:')
bullet(' cuanto más avanzás, más arriesgar conviene. El comodín, el cruce y la definición por penales son tus herramientas para dar el salto.','Arriesgá en los playoffs:')
bullet(' en tarjetas, penales y cambios tenés que acertar el número justo. Si no, ese ítem es cero.','Acertá exacto:')
bullet(' guardalo para una llave que valga la pena (semis). En la final ya no se puede usar.','Usá bien el comodín:')

doc.save(OUT_DOCX); print('docx generado:', OUT_DOCX); print('link:', LINK)
